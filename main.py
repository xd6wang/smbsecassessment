# -*- coding: utf-8 -*-

import docx
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING, WD_LINE_SPACING
import json
import metainfo
import content
from datetime import datetime


PROWLER_OUTPUT_FILE = "../prower/output/prowler-output.json"
REPORT_PATH_PREFIX = "./output/安全配置检查报告"


class Report(object):
    def __init__(self, path):
        self.path = path
        self.doc = docx.Document()
        self.finding_cnt = 1

    def add_heading(self, t):
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(20)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para = paragraph.add_run(t)
        para.font.size = Pt(18)
        para.font.name = "Microsoft YaHei"
        para._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    def add_title(self, t):
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.alignment = 1
        para = paragraph.add_run(t)
        para.font.size = Pt(26)
        para.font.name = "Microsoft YaHei"
        para._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    def add_paragraph(
        self,
        t,
        indent_for_newline=True,
        space_before_pt=None,
        space_after_pt=None,
        font_size=12,
    ):
        paragraph = self.doc.add_paragraph()
        if space_before_pt is not None:
            paragraph.paragraph_format.space_before = Pt(space_before_pt)
        if space_after_pt is not None:
            paragraph.paragraph_format.space_after = Pt(space_after_pt)
        if indent_for_newline:
            para = paragraph.add_run(" " * 7 + t)
        else:
            para = paragraph.add_run(t)
        para.font.size = Pt(font_size)
        para.font.name = "Microsoft YaHei"
        para._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        para.alignment = 1

    def add_finding_table(self, values):
        self.add_paragraph(
            "问题 {}".format(self.finding_cnt),
            indent_for_newline=False,
            space_before_pt=10,
            space_after_pt=0,
            font_size=10,
        )
        self.finding_cnt += 1
        table = self.doc.add_table(rows=8, cols=2)
        table.style = "Table Grid"

        cell = table.cell(0, 0)
        cell.text = "检查名称"
        cell = table.cell(0, 1)
        cell.text = values[0]

        cell = table.cell(1, 0)
        cell.text = "安全维度"
        cell = table.cell(1, 1)
        cell.text = values[1]

        cell = table.cell(2, 0)
        cell.text = "检查编号"
        cell = table.cell(2, 1)
        cell.text = values[2]

        cell = table.cell(3, 0)
        cell.text = "资源类型"
        cell = table.cell(3, 1)
        cell.text = values[3]

        cell = table.cell(4, 0)
        cell.text = "威胁描述"
        cell = table.cell(4, 1)
        cell.text = values[4]

        cell = table.cell(5, 0)
        cell.text = "缓解措施"
        cell = table.cell(5, 1)
        cell.text = values[5]

        cell = table.cell(6, 0)
        cell.text = "参考文档"
        cell = table.cell(6, 1)
        cell.text = values[6]

        cell = table.cell(7, 0)
        cell.text = "检查结果"
        cell = table.cell(7, 1)
        cell.text = values[7]

        for row in table.rows:
            c = 0
            for cell in row.cells:
                if c == 0:
                    cell.width = Cm(2.0)
                else:
                    cell.width = Cm(13.45)
                paragraphs = cell.paragraphs
                paragraph = paragraphs[0]
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.space_after = Pt(0)
                run_obj = paragraph.runs
                run = run_obj[0]
                font = run.font
                font.size = Pt(10)
                font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                c = c + 1

    def save(self):
        self.doc.save(self.path)


def get_count_by_check_category(category):
    return len([x for x in metainfo.check_group_map.values() if x[0] == category])


def load_prowler_output():
    findings = []
    findings_by_check = {}
    regions = set()
    count_non_info = 0
    count_fail = 0
    count_iam_check = get_count_by_check_category("iam")
    count_infra_check = get_count_by_check_category("infra")
    count_detection_check = get_count_by_check_category("detection")
    count_data_check = get_count_by_check_category("data")
    extra799_no_pass = True
    failed_checks_by_category = {
        "iam": set(),
        "infra": set(),
        "detection": set(),
        "data": set(),
    }
    failed_checks_by_subcategory = {
        "credentialleakage": set(),
        "leastpriviledge": set(),
        "publicaccess": set(),
        "audit": set(),
        "monitoring": set(),
        "encryption": set(),
    }

    with open(PROWLER_OUTPUT_FILE, "r") as f:
        lines = f.readlines()

        for line in lines:
            finding = json.loads(line)
            findings.append(finding)
            check_id = finding["Control"].split("]")[0][1:]
            regions.add(finding["Region"])
            account = finding["Account Number"]
            if check_id not in findings_by_check.keys():
                findings_by_check[check_id] = []
            if finding["Status"] == "FAIL":
                count_non_info += 1
                count_fail += 1
                findings_by_check[check_id].append(finding)
                category = metainfo.check_group_map[check_id][0]
                subcategory = metainfo.check_group_map[check_id][1]
                failed_checks_by_category[category].add(check_id)
                failed_checks_by_subcategory[subcategory].add(check_id)

            elif finding["Status"] == "PASS":
                count_non_info += 1
                if check_id == "extra799":
                    extra799_no_pass = False

        count_failed_checks = len([k for k, v in findings_by_check.items() if v])
        count_total_checks = len(findings_by_check.keys())

    return (
        findings_by_check,
        account,
        failed_checks_by_category,
        failed_checks_by_subcategory,
        len(regions),  # 0
        count_non_info,  # 1
        count_fail,  # 2
        count_failed_checks,  # 3
        count_total_checks,  # 4
        count_iam_check,  # 5
        count_infra_check,  # 6
        count_detection_check,  # 7
        count_data_check,  # 8
        extra799_no_pass,  # 9
    )


def write_report():
    (
        findings_by_check,
        account,
        failed_checks_by_category,
        failed_checks_by_subcategory,
        *ret,
    ) = load_prowler_output()

    report_path = "{}-{}-{}.docx".format(
        REPORT_PATH_PREFIX, account, datetime.now().strftime("%Y%m%d-%H%M%S")
    )
    report = Report(report_path)
    report.add_title(content.title)
    report.add_paragraph(
        content.summary.format(account, ret[4], ret[3], ret[3] / ret[4])
    )

    # add IAM sections
    report.add_heading(content.iam_title)
    report.add_paragraph(content.iam_general)
    if len(failed_checks_by_category["iam"]) == 0:
        report.add_paragraph(content.iam_pass)
    else:
        if len(failed_checks_by_subcategory["credentialleakage"]) != 0:
            report.add_paragraph(content.iam_credential_leak)
        if len(failed_checks_by_subcategory["leastpriviledge"]) != 0:
            report.add_paragraph(content.iam_least_priv)

    # add Infra Protection sections
    report.add_heading(content.infra_title)
    report.add_paragraph(content.infra_general)
    if len(failed_checks_by_category["infra"]) == 0:
        report.add_paragraph(content.infra_pass)
    else:
        report.add_paragraph(content.infra_public_access)

    # add Monitor sections
    report.add_heading(content.monitor_title)
    report.add_paragraph(content.monitor_general)
    if len(findings_by_check["extra713"]) == 0:
        report.add_paragraph(content.monitor_extra713_fail)
    if ret[9]:  # extra799_no_pass
        report.add_paragraph(content.monitor_extra799_no_pass)

    # add Data Protection sections
    report.add_heading(content.data_title)
    report.add_paragraph(content.data_general)
    if len(failed_checks_by_subcategory["encryption"]) == 0:
        report.add_paragraph(content.data_encrypt_general + content.data_encrypt_pass)
    else:
        report.add_paragraph(content.data_encrypt_general + content.data_encrypt_fail)

    # add Findings
    for check_id, findings in findings_by_check.items():
        if findings:
            report.add_finding_table(
                [
                    findings[0]["Control"].split("]")[1].lstrip(),
                    metainfo.category_translate_map[
                        metainfo.check_group_map[check_id][0]
                    ],
                    check_id,
                    findings[0]["Service"],
                    findings[0]["Risk"],
                    findings[0]["Remediation"],
                    findings[0]["Doc link"],
                    "\n".join([x["Message"] for x in findings]),
                ]
            )

    report.save()


if __name__ == "__main__":
    write_report()
